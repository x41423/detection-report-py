import os
import logging
from docx import Document
import re


class DataTransferModel:
    def __init__(self):
        pass

    def get_table_info(self, file_path: str) -> dict:
        """
        获取Word文档中表格的基本信息。
        
        Args:
            file_path: Word文档路径
            
        Returns:
            包含表格信息的字典：{"table_count": 表格数量, "first_table_rows": 第一个表格行数, 
                               "first_table_cols": 第一个表格列数, "sample_data": 前几行数据示例}
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        doc = Document(file_path)
        if not doc.tables:
            return {"table_count": 0, "error": "文档中没有找到表格"}
        
        table_count = len(doc.tables)
        first_table = doc.tables[0]
        rows = len(first_table.rows)
        cols = len(first_table.columns)
        
        # 获取前3行数据作为示例（每行只显示前6列）
        sample_data = []
        for i, row in enumerate(first_table.rows[:3]):
            row_data = []
            for j, cell in enumerate(row.cells[:6]):  # 只显示前6列
                # 清理文本，去除多余空白
                text = cell.text.strip()
                # 替换多个空白为单个空格
                text = re.sub(r'\s+', ' ', text)
                row_data.append(text)
            sample_data.append(row_data)
        
        return {
            "table_count": table_count,
            "first_table_rows": rows,
            "first_table_cols": cols,
            "sample_data": sample_data
        }

    def extract_all_varieties(self, table_paths: list) -> list:
        """
        从多个大表文件中提取所有去重后的品种名。

        Args:
            table_paths: 大表文件路径列表

        Returns:
            去重后的品种名列表（按首次出现顺序）
        """
        variety_names = ["品种", "蔬菜品种", "名称", "菜名"]
        all_varieties = []
        seen = set()

        for table_path in table_paths:
            if not os.path.exists(table_path):
                continue
            try:
                doc = Document(table_path)
                if not doc.tables:
                    continue
                table = doc.tables[0]
                rows = len(table.rows)
                if rows < 2:
                    continue

                # 识别品种列
                variety_col = -1
                header_row = table.rows[0]
                for col_idx, cell in enumerate(header_row.cells):
                    cell_text = cell.text.strip()
                    for name in variety_names:
                        if name in cell_text:
                            variety_col = col_idx
                            break
                    if variety_col >= 0:
                        break

                if variety_col == -1:
                    variety_col = 1  # 默认第2列

                # 提取品种名
                for row_idx, row in enumerate(table.rows):
                    if row_idx == 0:
                        continue
                    if variety_col >= len(row.cells):
                        continue
                    variety_text = row.cells[variety_col].text.strip()
                    if variety_text and variety_text.lower() not in seen:
                        seen.add(variety_text.lower())
                        all_varieties.append(variety_text)
            except Exception as e:
                logging.error(f"读取品种名失败 {table_path}: {e}")
                continue

        logging.info(f"从大表中提取到 {len(all_varieties)} 个品种名")
        return all_varieties

    def extract_data_from_big_table(self, big_table_path: str, veg_names: list) -> list:
        """
        从大表中提取指定菜名的数据，只提取三列：品种、抑制%、结果。
        根据列标题自动识别列位置。
        
        Args:
            big_table_path: 大表Word文档路径
            veg_names: 菜名列表（用于匹配品种列）
            
        Returns:
            匹配的数据列表，每个元素是一个字典，格式：
            {"variety": "品种名", "rate": "抑制%", "result": "结果"}
        """
        if not os.path.exists(big_table_path):
            raise FileNotFoundError(f"大表文件不存在: {big_table_path}")
        
        doc = Document(big_table_path)
        if not doc.tables:
            raise ValueError("大表中未找到表格，请检查文件格式。")
        
        # 假设第一个表格包含数据
        table = doc.tables[0]
        matched_data = []
        
        # 记录表格基本信息
        rows = len(table.rows)
        cols = len(table.columns)
        logging.info(f"开始提取数据: 表格有 {rows} 行, {cols} 列")
        logging.info(f"要匹配的菜名: {veg_names}")
        
        # 列名变体定义
        variety_names = ["品种", "蔬菜品种", "名称", "菜名"]
        rate_names = ["抑制%", "抑制率", "抑制 %", "抑制率%", "抑制率（%）", "抑制率(%)"]
        result_names = ["结果", "检测结果", "判定结果", "结论"]
        
        # 识别标题行（第一行）中三列的索引
        variety_col = -1
        rate_col = -1
        result_col = -1
        
        if rows > 0:
            header_row = table.rows[0]
            for col_idx, cell in enumerate(header_row.cells):
                cell_text = cell.text.strip()
                cell_text_lower = cell_text.lower()
                
                # 匹配品种列
                if variety_col == -1:
                    for name in variety_names:
                        if name in cell_text or cell_text_lower == name.lower():
                            variety_col = col_idx
                            logging.info(f"识别到品种列: 第{col_idx}列，标题: '{cell_text}'")
                            break
                
                # 匹配抑制%列
                if rate_col == -1:
                    for name in rate_names:
                        if name in cell_text or cell_text_lower == name.lower():
                            rate_col = col_idx
                            logging.info(f"识别到抑制%列: 第{col_idx}列，标题: '{cell_text}'")
                            break
                
                # 匹配结果列
                if result_col == -1:
                    for name in result_names:
                        if name in cell_text or cell_text_lower == name.lower():
                            result_col = col_idx
                            logging.info(f"识别到结果列: 第{col_idx}列，标题: '{cell_text}'")
                            break
        
        # 如果未找到某列，使用默认列索引（品种第2列，抑制%第3列，结果第4列）
        if variety_col == -1:
            variety_col = 1  # 第2列（索引1）
            logging.warning(f"未找到品种列标题，使用默认列索引: {variety_col}")
        if rate_col == -1:
            rate_col = 2  # 第3列（索引2）
            logging.warning(f"未找到抑制%列标题，使用默认列索引: {rate_col}")
        if result_col == -1:
            result_col = 3  # 第4列（索引3）
            logging.warning(f"未找到结果列标题，使用默认列索引: {result_col}")
        
        logging.info(f"最终列索引 - 品种: {variety_col}, 抑制%: {rate_col}, 结果: {result_col}")
        
        # 预处理菜名列表：去除空格，转为小写用于比较
        veg_names_processed = [name.strip().lower() for name in veg_names if name.strip()]
        
        # 遍历表格行（跳过标题行，第一行是标题）
        for row_idx, row in enumerate(table.rows):
            if row_idx == 0:  # 跳过标题行
                continue
            
            # 检查品种列是否匹配
            if variety_col >= len(row.cells):
                continue  # 该行列数不足
            
            variety_cell = row.cells[variety_col]
            variety_text = variety_cell.text.strip()
            variety_text_lower = variety_text.lower()
            
            # 检查是否匹配任何菜名
            matched = False
            for veg_name in veg_names_processed:
                if variety_text_lower == veg_name or variety_text.strip() == veg_name.strip():
                    matched = True
                    break
            
            if not matched:
                continue
            
            # 提取三列数据
            rate_text = ""
            result_text = ""
            
            if rate_col < len(row.cells):
                rate_text = row.cells[rate_col].text.strip()
            if result_col < len(row.cells):
                result_text = row.cells[result_col].text.strip()
            
            row_data = {
                "variety": variety_text,
                "rate": rate_text,
                "result": result_text
            }
            matched_data.append(row_data)
            logging.info(f"第{row_idx}行匹配菜名: '{variety_text}', 抑制%: '{rate_text}', 结果: '{result_text}'")
        
        if not matched_data:
            logging.warning(f"未找到匹配的菜名: {veg_names}")
            # 添加调试信息：显示前几行的品种列内容
            logging.info("调试信息 - 前5行的品种列内容:")
            for i, row in enumerate(table.rows[:5]):
                if i == 0:  # 跳过标题行
                    continue
                if variety_col < len(row.cells):
                    cell_text = row.cells[variety_col].text.strip()
                    if cell_text:
                        logging.info(f"  第{i}行品种列: '{cell_text}'")
        
        logging.info(f"提取完成: 找到 {len(matched_data)} 条匹配数据")
        return matched_data

    def write_data_to_small_table(self, small_template_path: str, data: list, output_path: str):
        """
        将数据写入小表模板，只写入四列：编号、品种、抑制%、结果。
        
        Args:
            small_template_path: 小表模板Word文档路径
            data: 数据列表，每个元素是字典，格式：{"variety": "品种", "rate": "抑制%", "result": "结果"}
            output_path: 输出文件路径
        """
        if not os.path.exists(small_template_path):
            raise FileNotFoundError(f"小表模板文件不存在: {small_template_path}")
        
        if not data:
            raise ValueError("没有数据可写入")
        
        doc = Document(small_template_path)
        if not doc.tables:
            raise ValueError("小表模板中未找到表格，请检查文件格式。")
        
        # 假设第一个表格是目标
        table = doc.tables[0]
        
        # 检查表格列数
        num_cols = len(table.columns)
        if num_cols < 4:
            raise ValueError(f"小表模板列数不足（{num_cols}列），至少需要4列（编号、品种、抑制%、结果）。")
        
        # 从第二行开始写入（第一行可能是标题）
        start_row = 1
        for idx, row_data in enumerate(data):
            row_index = start_row + idx
            # 如果行不够，添加新行
            if row_index >= len(table.rows):
                table.add_row()
            
            # 写入编号（第1列）
            self._safe_write_cell(table.rows[row_index].cells[0], str(idx + 1))
            
            # 写入品种（第2列）
            variety = row_data.get("variety", "")
            self._safe_write_cell(table.rows[row_index].cells[1], variety)
            
            # 写入抑制%（第3列）
            rate = row_data.get("rate", "")
            self._safe_write_cell(table.rows[row_index].cells[2], rate)
            
            # 写入结果（第4列）
            result = row_data.get("result", "")
            self._safe_write_cell(table.rows[row_index].cells[3], result)
            
            # 其他列保持不变（不覆盖）
        
        # 保存文档
        doc.save(output_path)
        logging.info(f"数据已写入小表: {output_path}")

    def _safe_write_cell(self, cell, text: str):
        """安全地写入单元格，保留原有格式"""
        if not cell.paragraphs:
            cell.add_paragraph(str(text))
            return
        para = cell.paragraphs[0]
        if not para.runs:
            para.add_run(str(text))
            return
        para.runs[0].text = str(text)
        for i in range(1, len(para.runs)):
            para.runs[i].text = ""

    def process_transfer(self, big_table_path: str, small_template_path: str, 
                         veg_names: list, output_path: str) -> dict:
        """
        执行完整的数据迁移流程。
        
        Returns:
            包含结果统计的字典，例如：
            {"matched_count": 5, "written_count": 5, "output_file": "路径", "details": [...]}
        """
        # 提取数据
        matched_data = self.extract_data_from_big_table(big_table_path, veg_names)
        
        if not matched_data:
            return {
                "matched_count": 0,
                "written_count": 0,
                "output_file": None,
                "message": "未找到任何匹配的菜名",
                "details": []
            }
        
        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        # 写入小表
        self.write_data_to_small_table(small_template_path, matched_data, output_path)
        
        # 准备详细信息
        details = []
        for item in matched_data:
            detail = {
                "品种": item.get("variety", ""),
                "抑制%": item.get("rate", ""),
                "结果": item.get("result", "")
            }
            details.append(detail)
        
        logging.info(f"成功提取 {len(matched_data)} 条数据，匹配详情: {details}")
        
        return {
            "matched_count": len(matched_data),
            "written_count": len(matched_data),
            "output_file": output_path,
            "message": f"成功提取 {len(matched_data)} 条数据并写入小表",
            "details": details
        }

    def process_multiple_tables(self, table_paths: list, small_template_path: str,
                                veg_names: list, output_path: str) -> dict:
        """
        处理多个大表文件，提取匹配数据并写入小表。
        
        Args:
            table_paths: 大表文件路径列表
            small_template_path: 小表模板路径
            veg_names: 菜名列表
            output_path: 输出文件路径
            
        Returns:
            包含结果统计的字典
        """
        if not table_paths:
            return {
                "processed_files": 0,
                "matched_count": 0,
                "written_count": 0,
                "output_file": None,
                "message": "未提供大表文件",
                "details": []
            }
        
        all_matched_data = []
        processed_files = 0
        
        # 按顺序处理每个大表文件
        for table_path in table_paths:
            try:
                logging.info(f"处理大表文件: {table_path}")
                matched_data = self.extract_data_from_big_table(table_path, veg_names)
                if matched_data:
                    all_matched_data.extend(matched_data)
                    processed_files += 1
                    logging.info(f"从 {os.path.basename(table_path)} 提取到 {len(matched_data)} 条数据")
                else:
                    logging.info(f"从 {os.path.basename(table_path)} 未提取到数据")
            except Exception as e:
                logging.error(f"处理文件 {table_path} 时出错: {e}")
                # 继续处理其他文件
                continue
        
        if not all_matched_data:
            return {
                "processed_files": processed_files,
                "matched_count": 0,
                "written_count": 0,
                "output_file": None,
                "message": "未找到任何匹配的菜名",
                "details": []
            }
        
        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        # 写入小表
        self.write_data_to_small_table(small_template_path, all_matched_data, output_path)
        
        # 准备详细信息
        details = []
        for item in all_matched_data:
            detail = {
                "品种": item.get("variety", ""),
                "抑制%": item.get("rate", ""),
                "结果": item.get("result", "")
            }
            details.append(detail)
        
        logging.info(f"共处理 {processed_files} 个大表文件，提取 {len(all_matched_data)} 条数据")
        
        return {
            "processed_files": processed_files,
            "matched_count": len(all_matched_data),
            "written_count": len(all_matched_data),
            "output_file": output_path,
            "message": f"成功从 {processed_files} 个大表文件提取 {len(all_matched_data)} 条数据并写入小表",
            "details": details
        }