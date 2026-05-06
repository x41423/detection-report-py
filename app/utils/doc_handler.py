import os
import re
from docx import Document


def replace_by_indices(para, start_idx: int, end_idx: int, new_str: str):
    """精准替换指定位置的字符，保持Word段落的原始排版。"""
    char_map = []
    for r_idx, run in enumerate(para.runs):
        for c_idx, char in enumerate(run.text):
            char_map.append((r_idx, c_idx))

    if not char_map or len(char_map) < end_idx:
        return

    run_texts = [list(run.text) for run in para.runs]

    for i in range(start_idx, end_idx):
        r_idx, c_idx = char_map[i]
        run_texts[r_idx][c_idx] = ""

    first_r_idx, first_c_idx = char_map[start_idx]
    run_texts[first_r_idx][first_c_idx] = new_str

    for r_idx, run in enumerate(para.runs):
        run.text = "".join(run_texts[r_idx])


def update_date_in_paragraph(para, new_date_str: str):
    """更新段落中的检测日期。"""
    original_text = para.text
    new_text = re.sub(r'检测日期：\s*(.+?)(?=\s*(主检|检测人|$))', f'检测日期：{new_date_str}', original_text)
    if new_text != original_text:
        para.text = new_text
        return True
    return False


def update_inspector_in_paragraph(para, inspector_name: str):
    """更新段落中的主检/检测人。"""
    original_text = para.text
    new_text = re.sub(r'(主检|检测人)：\s*(.+?)(?=\s*(检测日期|$))', f'\\1：{inspector_name}', original_text)
    if new_text != original_text:
        para.text = new_text
        return True
    return False


def modify_doc_absolutely_safe(doc: Document, new_date_str: str, inspector_name: str = "朱林初"):
    """修改检测日期和主检/检测人字段，保持排版不变。"""
    for para in doc.paragraphs:
        if "检测日期：" in para.text or "主检：" in para.text or "检测人：" in para.text:
            text = para.text
            
            # 匹配检测日期字段，直到「主检」「检测人」或段尾
            match_date = re.search(r'检测日期：(.+?)(?=\s*(?:主检|检测人|$))', text)
            if match_date and match_date.group(1).strip() != new_date_str:
                new_date_field = f"检测日期：{new_date_str}"
                replace_by_indices(para, match_date.start(0), match_date.end(0), new_date_field)
                # 重新获取文本，因为替换后文本改变
                text = para.text
            
            # 匹配主检/检测人字段，直到「检测日期」或段尾
            match_inspector = re.search(r'(主检|检测人)：(.+?)(?=\s*(?:检测日期|$))', text)
            if match_inspector and match_inspector.group(2).strip() != inspector_name:
                new_inspector_field = f"{match_inspector.group(1)}：{inspector_name}"
                replace_by_indices(para, match_inspector.start(0), match_inspector.end(0), new_inspector_field)


def safe_write_cell(cell, text: str):
    """在单元格内安全写入文本，保留原有格式结构。"""
    if not cell.paragraphs:
        cell.add_paragraph(str(text))
        return
    para = cell.paragraphs[0]
    if not para.runs:
        para.add_run(str(text))
        return
    para.runs[0].text = str(text)
    for i in range(1, len(para.runs)):
        para.runs[i].text = ""


def fill_table_data(table, data: list[dict], is_small: bool):
    """将数据填充到表格中。"""
    num_columns = len(table.columns)
    for i, item in enumerate(data):
        if not is_small:
            # 大表：前28条填入第一张表，其余填入第二张表。
            # 假设表格第一行是标题，数据从第2行开始（索引1）。
            row_index = i + 1
            if row_index >= len(table.rows):
                table.add_row()
            variety_col, rate_col, status_col = 1, 2, 3
        else:
            if i < 28:
                row_index, variety_col, rate_col, status_col = i + 1, 1, 2, 3
            elif num_columns >= 8:
                row_index, variety_col, rate_col, status_col = (i - 28) + 1, 5, 6, 7
            else:
                # 小表列数不足时，剩余数据不再填充
                break

        if row_index < len(table.rows):
            safe_write_cell(table.rows[row_index].cells[variety_col], item.get('variety', ''))
            safe_write_cell(table.rows[row_index].cells[rate_col], item.get('rate', ''))
            safe_write_cell(table.rows[row_index].cells[status_col], "合格")


def find_bx_table(doc: Document):
    """定位包含“滨鲜”的表格，用于小表模板。"""
    found = False
    for block in doc.element.body:
        if block.tag.endswith('p'):
            if "滨鲜" in "".join(block.itertext()):
                found = True
        elif block.tag.endswith('tbl') and found:
            for t in doc.tables:
                if t._element == block:
                    return t
    return doc.tables[0]


def process_single_document(doc_path: str, data_subset: list[dict], date_label: str, output_dir: str, inspector_name: str, is_small: bool):
    """处理单个文档：加载、修改、填充数据并保存。"""
    doc = Document(doc_path)
    if not doc.tables:
        raise ValueError("模板文档中未找到表格，请检查模板文件。")
    modify_doc_absolutely_safe(doc, date_label, inspector_name)
    fill_table_data(doc.tables[0], data_subset, is_small)
    output_path = os.path.join(output_dir, os.path.basename(doc_path))
    doc.save(output_path)


def process_documents(
    big_path: str,
    small_path: str,
    data: list[dict],
    date_label: str,
    output_dir: str,
    inspector_name: str = "朱林初",
):
    """处理大表和小表文档并保存到输出目录。"""
    if not os.path.exists(big_path) or not os.path.exists(small_path):
        raise FileNotFoundError("大表或小表文件不存在")

    if not os.path.isdir(output_dir):
        raise FileNotFoundError("输出目录不存在")

    if not data:
        raise ValueError("数据为空，无法生成报告。")

    # 小表容量校验：若列数不足 8，只支持 28 条
    doc_small_check = Document(small_path)
    small_table = doc_small_check.tables[0] if doc_small_check.tables else None
    if not small_table:
        raise ValueError("小表模板中未找到表格，请检查模板文件。")
    if len(small_table.columns) < 8 and len(data) > 28:
        raise ValueError("小表模板只有 4 列，仅支持 28 条数据，请分批生成或换 8 列模板。")

    # 大表分页：每 28 条生成一个文件，超过 1 页自动生成 -1, -2...
    page_size = 28
    total_pages = (len(data) + page_size - 1) // page_size
    for page_index in range(total_pages):
        start = page_index * page_size
        end = start + page_size
        page_data = data[start:end]

        if page_index == 0:
            output_name = os.path.basename(big_path)
        else:
            output_name = os.path.basename(big_path).replace(".docx", f"-{page_index}.docx")

        doc_b = Document(big_path)
        modify_doc_absolutely_safe(doc_b, date_label, inspector_name)
        fill_table_data(doc_b.tables[0], page_data, False)
        doc_b.save(os.path.join(output_dir, output_name))

    # 处理小表：所有数据填入小表（前28在左表，后28在右表，如果有8列）
    process_single_document(small_path, data, date_label, output_dir, inspector_name, is_small=True)
