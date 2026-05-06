import logging
import os
import re
from datetime import date
from docx import Document


_REPORT_DATE_RE = re.compile(
    r"农残检测记录表(?P<year>\d{4})[.年-](?P<month>\d{1,2})[.月-](?P<day>\d{1,2})日?(?:-(?P<page>\d+))?\.docx$",
    re.IGNORECASE,
)


class DocService:
    """Word 文档处理服务，迁移自 data_transfer_model.py"""

    def extract_all_varieties(self, table_paths: list[str]) -> list[str]:
        """从多个大表文件中提取所有去重后的品种名"""
        variety_names = ["品种", "蔬菜品种", "名称", "菜名"]
        all_varieties = []
        seen = set()

        for table_path in table_paths:
            if not os.path.exists(table_path):
                continue
            try:
                doc = Document(table_path)
                if not doc.tables:
                    continue
                table = doc.tables[0]
                if len(table.rows) < 2:
                    continue

                variety_col = self._find_column(table.rows[0], variety_names, default=1)

                for row_idx, row in enumerate(table.rows):
                    if row_idx == 0:
                        continue
                    if variety_col >= len(row.cells):
                        continue
                    text = row.cells[variety_col].text.strip()
                    if text and text.lower() not in seen:
                        seen.add(text.lower())
                        all_varieties.append(text)
            except Exception as e:
                logging.error(f"读取品种名失败 {table_path}: {e}")
                continue

        logging.info(f"从大表中提取到 {len(all_varieties)} 个品种名")
        return all_varieties

    def extract_data(self, big_table_path: str, veg_names: list[str]) -> list[dict]:
        """从大表中提取指定菜名的数据（品种、抑制%、结果）"""
        if not os.path.exists(big_table_path):
            raise FileNotFoundError(f"大表文件不存在: {big_table_path}")

        doc = Document(big_table_path)
        if not doc.tables:
            raise ValueError("大表中未找到表格")

        table = doc.tables[0]
        variety_names = ["品种", "蔬菜品种", "名称", "菜名"]
        rate_names = ["抑制%", "抑制率", "抑制 %", "抑制率%", "抑制率（%）", "抑制率(%)"]
        result_names = ["结果", "检测结果", "判定结果", "结论"]

        variety_col = self._find_column(table.rows[0], variety_names, default=1)
        rate_col = self._find_column(table.rows[0], rate_names, default=2)
        result_col = self._find_column(table.rows[0], result_names, default=3)

        veg_lower = [v.strip().lower() for v in veg_names if v.strip()]
        matched_data = []

        for row_idx, row in enumerate(table.rows):
            if row_idx == 0:
                continue
            if variety_col >= len(row.cells):
                continue

            variety_text = row.cells[variety_col].text.strip()
            if not variety_text:
                continue

            if variety_text.lower() not in veg_lower:
                continue

            rate = row.cells[rate_col].text.strip() if rate_col < len(row.cells) else ""
            result = row.cells[result_col].text.strip() if result_col < len(row.cells) else ""
            matched_data.append({"variety": variety_text, "rate": rate, "result": result})

        logging.info(f"从 {os.path.basename(big_table_path)} 提取到 {len(matched_data)} 条数据")
        return matched_data

    def write_to_small_table(self, template_path: str, data: list[dict], output_path: str):
        """将数据写入小表模板"""
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"小表模板文件不存在: {template_path}")
        if not data:
            raise ValueError("没有数据可写入")

        doc = Document(template_path)
        if not doc.tables:
            raise ValueError("小表模板中未找到表格")

        table = doc.tables[0]
        if len(table.columns) < 4:
            raise ValueError(f"小表模板列数不足（{len(table.columns)}列），至少需要4列")

        for idx, row_data in enumerate(data):
            row_index = 1 + idx
            if row_index >= len(table.rows):
                table.add_row()

            self._safe_write_cell(table.rows[row_index].cells[0], str(idx + 1))
            self._safe_write_cell(table.rows[row_index].cells[1], row_data.get("variety", ""))
            self._safe_write_cell(table.rows[row_index].cells[2], row_data.get("rate", ""))
            self._safe_write_cell(table.rows[row_index].cells[3], row_data.get("result", ""))

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        logging.info(f"数据已写入小表: {output_path}")

    def process_multiple_tables(self, table_paths: list[str], small_template_path: str,
                                 veg_names: list[str], output_path: str) -> dict:
        """处理多个大表文件，提取匹配数据并写入小表"""
        all_matched = []
        processed = 0

        for path in table_paths:
            try:
                matched = self.extract_data(path, veg_names)
                if matched:
                    all_matched.extend(matched)
                    processed += 1
            except Exception as e:
                logging.error(f"处理文件 {path} 失败: {e}")
                continue

        if not all_matched:
            return {
                "processed_files": processed, "matched_count": 0, "written_count": 0,
                "output_file": None, "message": "未找到任何匹配的菜名", "details": [],
            }

        # Deduplicate across multiple big tables by variety name (case-insensitive),
        # keeping the first match encountered.
        seen: set[str] = set()
        unique_matched: list[dict] = []
        for d in all_matched:
            key = str(d.get("variety", "")).strip().lower()
            if key and key not in seen:
                seen.add(key)
                unique_matched.append(d)
        all_matched = unique_matched

        self.write_to_small_table(small_template_path, all_matched, output_path)

        return {
            "processed_files": processed,
            "matched_count": len(all_matched),
            "written_count": len(all_matched),
            "output_file": output_path,
            "message": f"成功从 {processed} 个大表文件提取 {len(all_matched)} 条数据",
            "details": [{"variety": d["variety"], "rate": d["rate"], "result": d["result"]} for d in all_matched],
        }

    def preview_monthly_groups(self, table_paths: list[str], month: str) -> dict:
        grouped: dict[str, list[str]] = {}
        unrecognized: list[str] = []
        for path in table_paths:
            parsed_date = self.extract_report_date(path)
            if not parsed_date:
                unrecognized.append(os.path.basename(path))
                continue
            if parsed_date[:7] != month:
                unrecognized.append(os.path.basename(path))
                continue
            grouped.setdefault(parsed_date, []).append(path)

        groups = [
            {
                "date": date_text,
                "files": sorted(paths, key=self._report_page_sort_key),
                "count": len(paths),
            }
            for date_text, paths in sorted(grouped.items())
        ]
        return {
            "groups": groups,
            "unrecognized_files": unrecognized,
            "total_files": len(table_paths),
            "message": f"已识别 {len(groups)} 天、{sum(item['count'] for item in groups)} 个大表文件",
        }

    @staticmethod
    def extract_report_date(filename: str) -> str | None:
        match = _REPORT_DATE_RE.search(os.path.basename(str(filename)))
        if not match:
            return None
        year = int(match.group("year"))
        month = int(match.group("month"))
        day = int(match.group("day"))
        try:
            return date(year, month, day).isoformat()
        except ValueError:
            return None

    @staticmethod
    def _report_page_sort_key(path: str) -> tuple[int, str]:
        filename = os.path.basename(str(path))
        match = _REPORT_DATE_RE.search(filename)
        page = int(match.group("page") or 0) if match else 9999
        return page, filename

    @staticmethod
    def _find_column(header_row, candidates: list[str], default: int) -> int:
        """在标题行中查找匹配列"""
        for col_idx, cell in enumerate(header_row.cells):
            cell_text = cell.text.strip()
            for name in candidates:
                if name in cell_text:
                    return col_idx
        return default

    @staticmethod
    def _safe_write_cell(cell, text: str):
        """安全写入单元格，保留格式"""
        if not cell.paragraphs:
            cell.add_paragraph(str(text))
            return
        para = cell.paragraphs[0]
        if not para.runs:
            para.add_run(str(text))
            return
        para.runs[0].text = str(text)
        for i in range(1, len(para.runs)):
            para.runs[i].text = ""
